"""Build real Office/PDF documents from title + (markdown-ish) text.

Produces genuine binary files:
- ``.docx`` via python-docx
- ``.pdf``  via fpdf2

Libraries are imported lazily so a missing dependency degrades to a clear error
instead of breaking the backend. A small subset of markdown is supported:
``# heading``, ``## subheading``, and ``- ``/``* `` bullets; blank lines split
paragraphs.
"""

from __future__ import annotations

import io

from ..core.errors import AppError


def build_docx(title: str, content: str) -> bytes:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        raise AppError(
            "Word export needs python-docx. Install it with: pip install python-docx",
            status_code=503,
            code="docx_unavailable",
        ) from exc

    doc = Document()
    if title.strip():
        doc.add_heading(title.strip(), level=0)
    for raw in (content or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith(("- ", "* ")):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf(title: str, content: str) -> bytes:
    try:
        from fpdf import FPDF
    except Exception as exc:  # noqa: BLE001
        raise AppError(
            "PDF export needs fpdf2. Install it with: pip install fpdf2",
            status_code=503,
            code="pdf_unavailable",
        ) from exc

    # The built-in PDF fonts are latin-1; replace anything outside it so we never
    # crash (Spanish accents are covered; rare glyphs become '?').
    def safe(text: str) -> str:
        return text.encode("latin-1", "replace").decode("latin-1")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    if title.strip():
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 10, safe(title.strip()))
        pdf.ln(2)
    pdf.set_font("Helvetica", "", 12)
    for raw in (content or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            pdf.ln(4)
            continue
        if line.startswith("## "):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, safe(line[3:]))
            pdf.set_font("Helvetica", "", 12)
        elif line.startswith("# "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, safe(line[2:]))
            pdf.set_font("Helvetica", "", 12)
        elif line.startswith(("- ", "* ")):
            pdf.multi_cell(0, 7, safe("• " + line[2:]))
        else:
            pdf.multi_cell(0, 7, safe(line))
    return bytes(pdf.output())
